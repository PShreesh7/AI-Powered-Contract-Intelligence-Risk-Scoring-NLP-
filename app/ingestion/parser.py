"""
Day 1 module: document ingestion.
Extracts raw text from PDF and DOCX contract files.
"""
from __future__ import annotations
import os
import fitz  # PyMuPDF
import docx  # python-docx

MIN_TEXT_CHARS_PER_PAGE = 20

class UnsupportedFileTypeError(Exception):
    pass

def _ocr_pdf_page(file_path: str, page_number: int, dpi: int = 300) -> str:
    """Run Tesseract OCR on one PDF page using pdf2image."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError:
        # OCR dependencies aren't installed -- skip this page instead of
        # crashing the whole request. The rest of the document (any pages
        # with regular extractable text) will still process normally.
        return "" 

    images = convert_from_path(
        file_path,
        dpi=dpi,
        first_page=page_number,
        last_page=page_number,
    )

    if not images:
        return ""

    return pytesseract.image_to_string(images[0]).strip()

def extract_text_from_pdf(file_path: str) -> str:
    """Extract PDF text and fall back to OCR for scanned/image-only pages."""
    text_chunks = []

    with fitz.open(file_path) as doc:
        for page_number, page in enumerate(doc, start=1):
            page_text = page.get_text().strip()

            # If the PDF page has no useful text layer,
            # use OCR to extract text from the scanned image.
            if len(page_text) < MIN_TEXT_CHARS_PER_PAGE:
                page_text = _ocr_pdf_page(file_path, page_number)

            if page_text:
                text_chunks.append(page_text)

    return "\n".join(text_chunks).strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a Word document, including table cells."""
    document = docx.Document(file_path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


def extract_text(file_path: str) -> str:
    """Dispatch to the right extractor based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    # elif ext in (".docx", ".doc"):
    #     return extract_text_from_docx(file_path)

    elif ext == ".docx":
        return extract_text_from_docx(file_path)

    elif ext == ".doc":
        raise UnsupportedFileTypeError(
        "Legacy .doc files are not supported. "
        "Convert the file to .docx first."
    )
    else:
        raise UnsupportedFileTypeError(f"Unsupported file type: {ext}")


if __name__ == "__main__":
    import sys
    if len(sys.argv) != 2:
        print("Usage: python parser.py <path_to_contract>")
        sys.exit(1)
    print(extract_text(sys.argv[1])[:2000])
